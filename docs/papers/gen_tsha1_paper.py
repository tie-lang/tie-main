# -*- coding: utf-8 -*-
"""TSHA1 学术论文 docx 生成脚本（python-docx）。
排版：A4；正文宋体小四(12pt) 1.5 倍行距 + 首行缩进 2 字符；标题黑体三号居中；
英文 Abstract Times New Roman；表格 Table Grid。

描述对象：std/tsha1.tie 中重构后的 TSHA1 家族——f/r 复用「三进制双轨并行 + 最后综合」
压缩器；b 在双轨基础上增第三轨海绵轨（三轨并行）；tsha1x 独立四轨重构（双轨 + 海绵 + LFSR）。
本脚本内容与 std/tsha1.tie、std/base48.tie 及其探针生成器逐式对应。
"""
import os
import lxml.etree as _etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = os.environ.get("TSHA1_PAPER_OUT",
                     os.path.join(os.path.dirname(os.path.abspath(__file__)), "tsha1-paper.docx"))

doc = Document()

# ---- 页面：A4 + 页边距 ----
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.54)
sec.left_margin = sec.right_margin = Cm(3.17)

def _set_font(run, zh, en, size, bold=False, italic=False, color=None):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), zh)

def _norm(t):
    # 全角空格在某些字体/主题下会渲染为方框，统一替换半角
    return t.replace("\u3000", " ")

def para(text="", zh="宋体", en="Times New Roman", size=12, bold=False,
         align=None, indent=True, space_after=6, italic=False):
    text = _norm(text)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Pt(size * 2)   # 首行缩进 2 字符
    if align is not None:
        p.alignment = align
    if text:
        if "**" in text:
            # 支持 **加粗** 内联标记（python-docx 不解析 Markdown，手动分段渲染）
            for k, seg in enumerate(text.split("**")):
                if not seg:
                    continue
                r = p.add_run(seg)
                _set_font(r, zh, en, size, bold=(bold or (k % 2 == 1)), italic=italic)
        else:
            r = p.add_run(text)
            _set_font(r, zh, en, size, bold=bold, italic=italic)
    return p

def heading(text, size=14):
    text = _norm(text)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    r = p.add_run(text)
    _set_font(r, "黑体", "Times New Roman", size, bold=True)
    return p

def add_table(headers, rows, caption):
    caption = _norm(caption)
    para(caption, zh="黑体", size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=4)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        _set_font(r, "黑体", "Times New Roman", 10.5, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = ""
            r = cell.paragraphs[0].add_run(v)
            _set_font(r, "宋体", "Times New Roman", 10.5)
            if j == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para("", size=6, indent=False, space_after=4)
    return t

def _fmla(items, size=11, space_after=4):
    """渲染一行 Word 原生公式（OMML）。items: list[str] 即 OMML 片段，居中显示。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    pf.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    om = ('<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
          'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
          "<m:oMath>" + "".join(items) + "</m:oMath></m:oMathPara>")
    p._p.append(_etree.fromstring(om.encode("utf-8")))
    return p

# ---- OMML 构建小工具（均返回 m:oMath 内的子元素字符串）----
def _esc(text):
    """把普通文本转义为合法 XML 文本内容（公式中的 &、<、> 须转义）。"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def _t(text, zh=False):
    """数学文本运行（默认数学斜体字体；zh=True 时中文用宋体）。自动 XML 转义。"""
    rpr = '<w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr>' if not zh else \
          '<w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math" w:eastAsia="黑体"/></w:rPr>'
    return '<m:r>%s<m:t xml:space="preserve">%s</m:t></m:r>' % (rpr, _esc(text))

def _sub(base, sub):
    return '<m:sSub><m:e>%s</m:e><m:sub>%s</m:sub></m:sSub>' % (base, sub)

def _sup(base, sup):
    return '<m:sSup><m:e>%s</m:e><m:sup>%s</m:sup></m:sSup>' % (base, sup)

def _frac(num, den):
    return '<m:f><m:num>%s</m:num><m:den>%s</m:den></m:f>' % (num, den)

def _d(inner, beg="(", end=")"):
    """带括号定界符（自动伸缩括号）。"""
    return '<m:d><m:dPr><m:begChr m:val="%s"/><m:endChr m:val="%s"/></m:dPr><m:e>%s</m:e></m:d>' % (beg, end, inner)

def _nary(ch, sub, sup, e):
    """n 元算子（求和/连乘）：ch 为算子字符（∑ / ∏）。"""
    return '<m:nary><m:naryPr><m:chr m:val="%s"/><m:limLoc m:val="undOvr"/></m:naryPr>' \
           '<m:sub>%s</m:sub><m:sup>%s</m:sup><m:e>%s</m:e></m:nary>' % (ch, sub, sup, e)

def _arr(rows):
    """方程组/多行数组（每行一个 m:e）。"""
    return '<m:d><m:dPr><m:begChr m:val="["/><m:endChr m:val=""/></m:dPr>%s</m:d>' % \
           "".join("<m:e>%s</m:e>" % r for r in rows)

def _mat(rows):
    """矩阵/多行并列（写为方程数组 m:eqArr，无括号）。"""
    return '<m:eqArr>%s</m:eqArr>' % "".join("<m:e>%s</m:e>" % r for r in rows)

def _fhead(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    r = p.add_run(_norm(text))
    _set_font(r, "黑体", "Times New Roman", 11, bold=True)
    return p

# ================= Word 原生绘图（VML）辅助 =================
def _vtx(lines, sz, bold=False):
    """VML 文本框多行内容（w:txbxContent，黑体，sz 为半磅）。"""
    runs = ""
    for t in lines:
        b = "<w:b/>" if bold else ""
        runs += ('<w:p><w:pPr><w:spacing w:before="0" w:after="0"/><w:jc w:val="center"/></w:pPr>'
                 '<w:r><w:rPr><w:rFonts w:ascii="SimHei" w:hAnsi="SimHei" w:eastAsia="SimHei"/>'
                 '%s<w:sz w:val="%d"/><w:szCs w:val="%d"/></w:rPr>'
                 '<w:t xml:space="preserve">%s</w:t></w:r></w:p>' % (b, sz, sz, t))
    return ('<v:textbox inset="4,3,4,3" style="mso-fit-shape-to-text:t">'
            '<w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "%s</w:txbxContent></v:textbox>" % runs)

def _vrect(x, y, w, h, lines, fill, edge, sz=20, bold=False):
    return ('<v:rect style="position:absolute;left:%d;top:%d;width:%d;height:%d;v-text-anchor:middle" '
            'fillcolor="%s" strokecolor="%s" strokeweight="1.5pt">%s</v:rect>'
            % (x, y, w, h, fill, edge, _vtx(lines, sz, bold)))

def _vline_arrow(x1, y1, x2, y2, color="#33475b", wd="2.25pt", endarrow=True):
    s = '<v:line from="%d,%d" to="%d,%d" strokecolor="%s" strokeweight="%s">' % (x1, y1, x2, y2, color, wd)
    if endarrow:
        s += '<v:stroke endarrow="classic" endarrowwidth="medium" endarrowlength="medium"/>'
    s += "</v:line>"
    return s

def add_vml_arch():
    """插入 Word 原生绘图架构图（VML 形状 + 连接线箭头），与 PNG 版布局一致。"""
    shapes = []
    # 顶部输入与填充
    shapes.append(_vrect(570, 26, 340, 58, ["消息 M（任意字节串）"], "#fdf3dc", "#a8863d", sz=24))
    shapes.append(_vrect(528, 112, 394, 62, ["纯零填充 + 64 位计数器 t", "（mod 2^64，长度绑定）"], "#ffffff", "#33475b", sz=20))
    shapes.append(_vline_arrow(740, 84, 740, 112))
    shapes.append(_vline_arrow(740, 174, 740, 196))
    # 主线：从最左栏中心延伸到最右栏中心（修复左侧悬空断线）
    cx0, cxn = 228, 1272
    shapes.append(_vline_arrow(cx0, 196, cxn, 196, endarrow=False))
    # 四模型：f/r 双轨、b 三轨（含海绵）、x 四轨（含 LFSR），仅轮数/输出字数不同
    cols = [
        (76, 228, "tsha1f-256（快速模型）", ["三进制双轨并行 R_F=12", "trit 位平面（平衡加/乘/多数）", "双轨耦合 + 轮常量，末段 S=4 综合"], "#e8f4f8"),
        (424, 576, "tsha1b-256（复杂模型）", ["三轨并行：双轨 + 第三轨海绵 R_B=14", "海绵吸收 / 8 字置换 / 双轨耦合", "末段 S=4 综合"], "#f3eef9"),
        (772, 924, "tsha1x-256（加强模型）", ["四轨并行：双轨 + 海绵 + LFSR R_X=16", "LFSR 反馈⊕量化门控去线性化", "末段 fin_synth_x S=4 综合"], "#fdeeee"),
        (1120, 1272, "tsha1r-128（轻量模型）", ["三进制双轨并行 R_R=8", "trit 位平面扩散（轮数最少）", "末段 S=4 综合，仅取前 4 字"], "#eef9ef"),
    ]
    for x, cx, name, steps, fill in cols:
        shapes.append(_vline_arrow(cx, 196, cx, 246))
        shapes.append(_vrect(x, 260, 304, 58, [name], "#ffffff", "#33475b", sz=21, bold=True))
        y = 332
        for s in steps:
            shapes.append(_vrect(x, y, 304, 54, [s], fill, "#7d8ea3", sz=18))
            y += 66
    # 输出
    for _, cx, _, _, _ in cols:
        shapes.append(_vline_arrow(cx, 530, cx, 694, color="#7d8ea3"))
    shapes.append(_vrect(76, 700, 1348, 76,
                         ["任选一模型：位长 n 显式指定，n∈{2,3,4,6,8,12,16,24,32,48,64,88,96}（48 进制符号个数）",
                          "→ Base48 编码（n 个字符）｜ 其它进制 {2,3,8,16} 按信息量换算 ｜ _hex 可选（64/32）"],
                         "#ffffff", "#33475b", sz=19))
    body = "".join(shapes)
    vml = ('<w:pict xmlns:v="urn:schemas-microsoft-com:vml" '
           'xmlns:o="urn:schemas-microsoft-com:office:office" '
           'xmlns:w="urn:schemas-microsoft-com:office:word">'
           '<v:group style="width:411pt;height:232pt" coordsize="1500,850" coordorigin="0,0">'
           "%s</v:group></w:pict>" % body)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run._r.append(_etree.fromstring(vml.encode("utf-8")))

# ================= 标题 =================
para("TSHA1：一种多原语确定性组合的安全杂凑函数家族", zh="黑体", size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=2)
para("——统一三进制双轨并行构造、强度梯度与可论证安全性质（基于自举语言 TIE 的纯语言实现）", zh="黑体", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=10)
para("侯杨宝鑫，TIE 项目团队", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=2)
para("（侯杨宝鑫为 TIE 项目团队成员；单位：TIE Language Project）", zh="楷体", size=10.5,
     align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=10)

# ================= 中文摘要 =================
heading("摘　要", size=12)
para("TSHA1 是为自举编程语言 TIE 设计并完全以该语言实现的**新型安全杂凑函数家族**，"
     "包含快速模型 tsha1f、复杂模型 tsha1b、加强模型 tsha1x 与轻量模型 tsha1r。其“新型”体现在"
     "统一的内核构造：f/b/r 共享同一套平衡三进制位平面压缩原语（trit 位平面平衡加 "
     "tadd2、平衡乘 tmul2、majority 量化 quant3、旋转混洗 rrp，全部以 &/^/|/<< 位运算实现，"
     "无逐 trit 模 3 慢点），轨 A 与轨 B 各自四组三进制字并行 R 轮扩散，每轮消息三进制平面"
     "交替注入两轨并做轨间耦合与轮常量错位注入；f 取双轨并行 R_F=12、r 取双轨并行 R_R=8，"
     "b 在双轨基础上增第三轨海绵轨（消息吸收/8 字置换/与双轨双向耦合，R_B=14）；整条链压缩完后"
     "以 S=4 轮**最后综合**（fin_synth）收束并投影回链值，仅轮数与种子派生常量不同。"
     "tsha1x 则独立重构为**四轨并行**（双轨 + 海绵轨 + LFSR 反馈移位轨，R_X=16），以独立常量族"
     "并以 fin_synth_x 终筛，为签名对象/凭证提供最高强度。该家族以统一压缩原语构建四档速度/强度梯度，"
     "避免多骨架维护面。其安全性来自四类可论证性质：其一，计数器式纯零填充与 64 位长度绑定在"
     "结构上免疫 Merkle–Damgård 型长度扩展攻击；其二，全部常量（IV 与轮常量，含海绵轨/LFSR 轨"
     "独立常量）由“标准种子＋PRNG 扩展”即 SHA-256(SEED ‖ u64be(k)) 计数器流确定性生成、可复现，"
     "杜绝猜测常量与后门注入疑虑；其三，自创层严格限于参数、常量、轮数与并行结构，安全性质可随"
     "组合逐一论证；其四，三进制平衡加/乘与 majority 的非线性交叠，配合加强模型的多轨并行，共同"
     "抬高结构化攻击的破解成本。实现为纯 TIE、零外部依赖，逐**字节**处理任意 UTF-8 原始字节。"
     "摘要默认输出为自行设计的 Base48（48 符号连续字符集）编码，位长 n 显式指定且仅允许"
     "{2,3,4,6,8,12,16,24,32,48,64,88,96}（48 进制符号个数），不足时按 XOF 流扩展，便于"
     "人类抄写、指纹与审计链标识；同时保留 _hex 变体（内部摘要 64/32 hex）与 {2,3,8,16} "
     "进制换算。探针以跨平台交叉生成的已知答案向量逐字节核对。基准实测（TIE 自写编译器/IR "
     "执行层）大消息稳态吞吐 tsha1r≈42 MB/s、tsha1f≈33、tsha1b≈16、tsha1x≈11，模型间"
     "速度差与强度梯度设计相符。该家族已作为插件化内核审计链的指纹与凭证签名哈希底座投入使用。",
     size=12, space_after=6)
p = para("", size=12, indent=False, space_after=10)
r0 = p.add_run("关键词：")
_set_font(r0, "黑体", "Times New Roman", 12, bold=True)
for w in ["新型安全杂凑", "确定性组合构造", "三进制双轨并行", "平衡三进制", "抗长度扩展",
          "XOF", "强度梯度", "Base48"]:
    p.add_run(w + " ")

# ================= English Abstract =================
heading("Abstract", size=12)
para("TSHA1 is a new family of cryptographic hash functions designed for and fully implemented "
     "in the self-hosting programming language TIE, comprising four variants: tsha1f (fast), "
     "tsha1b (complex), tsha1x (strengthened) and tsha1r (lightweight). Its novelty lies in a "
     "unified construction rather than a single primitive: f/b/r share one set of balanced-ternary "
     "bit-plane primitives (balanced add tadd2, balanced multiply tmul2, majority quant3 and "
     "rotation mix rrp, all expressed in &/^/|/<< with no per-trit mod-3 bottleneck). Track A "
     "and Track B each hold four ternary words diffused in parallel, with cross-track coupling "
     "and offset round-constant injection; f runs the dual-track kernel for R_F=12 rounds and r "
     "for R_R=8, while b augments the two tracks with a third sponge track (rate/capacity; "
     "message absorption, 8-word permutation and bidirectional coupling with the two tracks) "
     "for R_B=14 rounds. After the whole chain is compressed, an S=4 final-synthesis round "
     "(fin_synth) folds the chain back into the hash, differing only in the number of rounds "
     "and the seeded constants. tsha1x is independently re-built as a four-track parallel "
     "kernel--dual track plus sponge track plus LFSR feedback-shift track, R_X=16--using its own "
     "seed-derived constant family and a fin_synth_x final synthesis, delivering the highest "
     "strength for signature objects. Improved security is argued on four verifiable grounds: "
     "(i) counter-based pure-zero padding with a "
     "64-bit length binding structurally defeats Merkle-Damgard-style length-extension attacks; "
     "(ii) all constants (IV and round constants) are deterministically derived from a standard "
     "seed via PRNG extension (SHA-256(SEED || u64be(k)) counter stream) and are reproducible, "
     "ruling out guess-constant or backdoor concerns; (iii) the self-designed layer is restricted "
     "to parameters/constants/rounds/permutation order, so security can be argued compositionally; "
     "and (iv) the interplay of balanced-ternary add/multiply/majority non-linearity, combined "
     "with the strengthened variant's multi-track parallel construction, raises the cost of structural "
     "cryptanalysis. The implementation is pure TIE with no external dependency, processing "
     "arbitrary UTF-8 bytes byte-by-byte; digests default to a purpose-built Base48 alphabet with "
     "an explicit bit-length n drawn from {2,3,4,6,8,12,16,24,32,48,64,88,96} (Base48 symbols), "
     "extended via an XOF stream when short, while _hex variants (64/32) and base {2,3,8,16} "
     "mappings are also provided. Known-answer vectors cross-generated with the platform are "
     "verified byte-for-byte. Benchmarks on the TIE self-hosted execution layer show "
     "approximately 42/33/16/11 MB/s for r/f/b/x on large messages, matching the intended "
     "hierarchy. TSHA1 has been adopted as the fingerprint and credential-signing hash "
     "foundation of the plug-in kernel audit chain.",
     size=12, space_after=6)
p = para("", size=12, indent=False, space_after=10)
r0 = p.add_run("Keywords: ")
_set_font(r0, "宋体", "Times New Roman", 12, italic=True)
for w in ["new design hash", "deterministic composition", "ternary dual-track parallel",
          "balanced ternary", "length-extension resistance", "XOF", "strength hierarchy", "Base48"]:
    p.add_run(w + " ")

# ================= 1 引言 =================
heading("1　引言")
para("哈希函数是现代密码体系的地基，其安全性直接决定签名、消息认证、密钥派生与完整性"
     "审计的可靠性。标准算法虽经长期沉淀而高度成熟，仍存在结构性局限：MD5/SHA-1 已出现"
     "可证明的碰撞攻击；SHA-2 的 Merkle–Damgård 结构公开了长度扩展攻击面；SHA-3 虽结构"
     "免疫，但其海绵构造以吞吐换取状态规模，在受限环境中开销不小。与此同时，越来越多"
     "自给自足的平台希望拥有“可自持、可复现、安全性质可明确论证”的自有哈希体系，而非"
     "仅仅封装外部标准实现。TIE 便是一个以自身编译器（tiec）构建的自举编程语言系统："
     "其前端、中端、后端与标准库全部以 TIE 编写，形成“tie 写 tiec、tie 写 tie 库”的闭环。"
     "在向全平台插件化架构演进的过程中，平台需要一个完全由 TIE 自身实现、随自举闭环一致"
     "交付的安全哈希底座，用于插件审计链的完整性指纹、包树根校验与凭证签名对象绑定。")
para("本文提出的 TSHA1（全称 TIE Secure Hash Algorithm 1.0，下文简称 TSHA1）正是面向"
     "“新型、更安全”这一目标设计的。"
     "其**新颖性**体现为统一内核构造：f/b/r 共享同一套 trit 位平面压缩原语（平衡加/平衡乘/"
     "majority/rrp，全位运算）与双轨并行骨架——轨 A 与轨 B 四组三进制字并行扩散 + 轨间耦合 + "
     "轮常量错位注入，唯一区别是轮数与由种子派生的常量；b 在双轨基础上增第三轨海绵轨；"
     "加强模型 tsha1x 则独立重构为四轨并行（双轨 + 海绵 + LFSR 反馈移位轨），以 R_X=16 轮"
     "提供全家族最高强度。"
     "这样既通过调节轮数与轨道数获得四档速度/强度梯度，又只需维护统一的压缩原语，避免多骨架的"
     "实现与验证面。其**更强的安全性质**可归纳为：(1) 计数器式纯零填充与 64 位长度绑定在"
     "结构上免疫长度扩展攻击；(2) 全部常量由“标准种子＋PRNG 扩展”确定性生成、可复现，"
     "排除猜测常量与后门注入疑虑；(3) 自创层严格限于参数、常量、轮数与并行结构，未引入"
     "未经审计的全新结构断言，安全性可组合论证；(4) 三进制平衡加/乘与 majority 非线性交叠，"
     "配合加强模型的多轨并行，整体抬高结构化攻击的破解成本；(5) 四模型梯度使同一族可"
     "同时满足快速指纹、复杂强度与签名对象级防护。trit 位平面扩散在这一构造中扮演核心扩散"
     "构件角色（§3.1–§3.2），其实现受益于 TIE 语言对平衡三进制的一等支持，但本文的安全主张"
     "不拘泥于三进制本身，而是重视其经位平面化后全部以 &/^/|/<< 实现的工程与安全收益。")
para("本文组织如下：第 2 节综述相关工作；第 3 节给出 TSHA1 设计与四模型结构；第 4 节讨论"
     "安全性质与已知限制；第 5 节报告纯 TIE 实现、向量验证与基准；第 6 节说明其在插件"
     "审计链中的应用；第 7 节总结与展望。")

para("现状与当前存在的问题。就本研究启动时而言：(1) 平台对标准哈希（SHA-2/3、BLAKE2/3）"
     "已有标准库实现，但作为外部标准算法，其常量与参数不体现平台语境，难以满足“自持有、"
     "且随自举闭环逐字节复现”的审计链强需求；(2) 标准算法各自存在结构性局限（长度扩展、"
     "温室化参数、退出审计较难的长期依赖），而评估一个“既有实力又可控”的自有族需要可论证"
     "的安全机制，而非经验式拼装；(3) 平台执行模型（有符号 64 位、表驱动、字符串 {ptr,len} "
     "逐字节语义）对通用算法的移植存在适配成本与验证噪音；(4) 插件化审计链需要“文件级快速"
     "指纹—包级强指纹—签名对象”的分层哈希底座，单一算法难以同时满足成本与强度梯度。")
para("针对上述现状，TSHA1 着力解决以下问题：（1）自持——以纯 TIE 实现四模型哈希，零外部"
     "运行时依赖，常量按“标准种子＋PRNG 扩展”可复现（§3.7），并随自举闭环交付（§5.1）；"
     "（2）安全机制可论证——抗长度扩展的结构设计（§3.1）、可复现常量（§3.7）、自创层限于"
     "参数/常量/轮数与单一内核的高可验证性（§4），使“更安全”不依赖对未验证结构的断言；"
     "（3）模型深度——四模型按速度/强度梯度组织，实测 r≈42 > f≈33 > b≈16 > x≈11 MB/s（§5.3），"
     "单调满足审计链 per-file / 包树根 / 签名对象的三级需求（§6）；（4）适配验证——全部"
     "实现遵循 i64 无溢出论证与逐字节模型，探针以跨平台交叉生成的 KAT 逐字节核对，并与"
     "常规回归体系解耦验证（§5.2）。")

# ================= 2 相关工作 =================
heading("2　相关工作")
para("现代杂凑函数的设计沿两条主线演进：(1) ARX（加-旋-异或）路线，以 BLAKE"
     "[3][4] 为代表，其 G 函数将模 2^32 加法、循环旋转与异或组合，轮数较少即可获得良好"
     "雪崩与扩散；BLAKE2（RFC 7693）[3] 与 BLAKE3 [5] 进一步在并行与速度上优化，"
     "BLAKE3 采用二叉树结构支持多核与 XOF 输出。(2) SPN（代换-置换网络）路线，以"
     "Keccak（SHA-3，FIPS 202）[2][11] 的海绵结构与 Ascon（NIST 轻量级加密标准，"
     "Feistel-SPN 混合）[6] 为代表，非线性特征清晰、侧信道设计友好。此外 SHA-2"
     "（FIPS 180-4）[1] 与 SHA-3 均由美国 NIST 标准化，是联邦与现代协议的事实标准。"
     "三进制/平衡三进制系统在进位链密码（如用平衡三进制实现进位无关的加乘混合扩散，"
     "以及 1 式 MVL 逻辑电路）中已有较久研究，TSHA1 将平衡三进制以位平面方式嵌入通用"
     "级联压缩内核，是其工程化应用的一个实例。")
para("在杂凑函数的验证与部署侧，已知答案向量（KAT）与测试向量生成器、NIST ACVP"
     "自动化验证、以及常数时间实现审查是通用工程实践。后量子时代 NIST 已发布 ML-KEM"
     "（FIPS 203）[7]、ML-DSA（FIPS 204）[12] 与 SLH-DSA（FIPS 205）[8]，其中 SLH-DSA"
     "以哈希基构建，与本文的“以自有哈希支撑平台”思路一致。TIE 平台的既有安全底座还"
     "包括 SHA-256/SHA-3/Keccak、BLAKE2、BLAKE3、Toom/分基大数（bigint）、Ed25519/X25519"
     "纯语言实现与各类 KDF/MAC 原语，TSHA1 与之在库中并存、互不替代：TSHA1 承担平台自有"
     "标准的、可随自举闭环交付的签名/指纹底座角色。")

# ================= 3 设计 =================
heading("3　TSHA1 设计")

heading("3.1　总体架构与四模型定位", size=12)
para("TSHA1 家族采用“同族多模型”策略：f/b/x/r 共享输出编码、验证探针与常量派生框架；"
     "f/b/r 共享**同一套 trit 位平面压缩原语与双轨骨架**，仅轮数与种子派生的 IV/轮常量（b 另增"
     "第三轨海绵轨）不同，从而以最小实现面权衡速度与强度（表 1）；tsha1x 则独立重构为四轨并行"
     "（双轨 + 海绵 + LFSR），以 R_X=16 提供签名对象级最强档（§3.5）。家族命名遵循 "
     "**tsha1[模型]**；对外接口为 tsha1f/tsha1b/tsha1x/"
     "tsha1r(msg, n, base)，其中位长 n 显式给出且**只能**取 "
     "{2,3,4,6,8,12,16,24,32,48,64,88,96}——它指输出摘要中 48 进制符号的个数（字符个数，"
     "不是比特数），例如 tsha1f(msg, 8) 恰好输出 8 个 48 进制字符；列表外的值（含 0、46、23）"
     "一律拒绝并返回空串，**没有默认位长概念**。可选参数 base∈{2,3,8,16,48} 指定输出进制，"
     "默认 48；n 始终是“48 进制下的位数”，其它进制按相同信息量换算（内嵌字节数换算见 §3.8）。"
     "对应内嵌字节数与 XOF 扩展流程详见 §3.8。")
add_table(
    ["模型", "内核/结构", "轮数（S 综合）", "内部摘要", "定位"],
    [
        ["tsha1f", "三进制双轨并行 + 最后综合", "R_F=12（S=4）", "32 B（64 hex）", "快速，海量小文件指纹"],
        ["tsha1b", "三轨并行[双轨+海绵] + 最后综合", "R_B=14（S=4）", "32 B（64 hex）", "复杂强度，审计链"],
        ["tsha1x", "四轨并行[双轨+海绵+LFSR] + 最后综合", "R_X=16（S=4）", "32 B（64 hex）", "加强，签名对象/凭证"],
        ["tsha1r", "三进制双轨并行 + 最后综合", "R_R=8（S=4）", "16 B（32 hex）", "嵌入式/超小消息"],
    ],
    "表 1　TSHA1 四模型定位（位长 n = 48 进制符号个数，仅允许 {2,3,4,6,8,12,16,24,32,48,64,88,96}）")

# 图 1：家族总体架构（Word 原生绘图对象，箭头由连接线实现，可正常显示）
add_vml_arch()
para("图 1　TSHA1 家族总体架构（f/r 复用三进制双轨并行内核；b 增第三轨海绵；tsha1x 四轨并行含 LFSR，均以 S=4 最后综合终筛）",
     zh="黑体", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=8)

heading("3.2　统一压缩原语与双轨骨架；b/x 的多轨扩展", size=12)
para("TSHA1 的全部模型建立在统一的平衡三进制位平面压缩原语之上，f/b/r 共享同一双轨并行骨架，"
     "仅轮数与常量（b 增第三轨海绵轨、x 独立四轨）不同。"
     "首先是 trit 位平面表示（层内必选设计）：每个 u64 装 32 个 trit，高 32 位为幅值位平面 "
     "M（第 i 位=1 当且仅当第 i 个 trit 非零），低 32 位为符号位平面 N（第 i 位=1 当且仅当第 i 个 trit 为负），"
     "即等价于 (M<<32)|N。TIE 以两个恒 <2^32 的 i64 半字 "
     "(M,N) 表示，因 M、N 均不含符号位，全部扩散运算可表达为 &、^、|、<<，避免逐 trit "
     "模 3 的慢点与条件分支。字节经饱和映射注入位平面：字节高二位 tf=(b>>6)&3 经 to_trit "
     "饱和（0→−1、1→0、2/3→+1）得 trit 值 tv∈{−1,0,+1}。")
para("双轨骨架以“三进制双轨并行”为基：工作状态 v[0..15] 分为轨 A（v[0..7]）与轨 B"
     "（v[8..15]），每两个位平面 (M,N) 构成一个三进制字，两轨各含 4 个三进制字。初始化时"
     "由链值 h 与 IV 经互补旋转载入双轨（轨 A←h、轨 B←rot(h)‖IV，再异或余下 IV 字），注入"
     "64 位计数器 t 的两半（t_lo/t_hi），末块再异或全 1 作终块标记；填充一律纯零（无 "
     "0x80/无长度位），长度由计数器承载。随后两轨各自独立跑 R 轮（f=12、r=8）三进制"
     "扩散：每轮由轮序 r 与 24 基调度键 skey 派生错位的移位量（rA/rB/mA/mB），轨内并行做"
     "平衡加 tadd2、平衡乘 tmul2 与 majority 量化 quant3（§A.3.1–A.3.2）；消息三进制平面 "
     "(M0,N0)/(M1,N1) 每轮交替注入两轨（轨 A←M0,N0；轨 B←M1,N1，旋转后平衡加）；随后做轨间"
     "耦合（隔轨平衡加 / 乘积 / 旋转传播）与轮常量双轨错位注入。b 在此双轨基础上增第三轨海绵轨"
     "（§3.4），x 独立重构为四轨并行（§3.5）。整条链压缩完后，末尾统一执行"
     "最后综合 fin_synth：把链值 h 重新载入双轨，跑 S=4 轮一体化收束轮（全 8 字环式平衡加 + "
     "乘积 + majority），再投影折叠回链值完成终筛。f/b/x 输出全部 8 字（64 hex）；r 仅取前 "
     "4 字（32 hex）。")

heading("3.3　tsha1f：快速模型（三进制双轨，R=12）", size=12)
para("tsha1f 是快速档，完整执行 3.2 节所描述的统一内核，取轮数 R_F=12（三档中最快）、"
     "常量族 IV_F/RCON_F（SEED_F=“TSHA1-2026-f-256-v1”派生）。其作用是提供海量小对象"
     "批量的廉价指纹，同时共享与 b 同构的内核实现与验证面。消息填充采用纯零填充（不含 "
     "0x80 与显式长度位），长度由 64 位计数器 t 的两半（t_lo、t_hi，各 32 位，模 2^64）承载；"
     "输出 8 字大端字节序、共 64 hex（内部摘要 32 字节）。逐字处理采用 str_byte + len，"
     "因此任意 UTF-8 原始字节序列均可哈希。")
_fhead("tsha1f 逐步计算式")
_fmla([_t("tadd2((Ma,Na),(Mb,Nb))：aP=Ma∧¬Na；aN=Ma∧Na；bP=Mb∧¬Nb；bN=Mb∧Nb")])
_fmla([_t("o_pos=(¬Ma∧bP)∨(aP∧¬Mb)∨(aN∧bN)；o_neg=(¬Ma∧bN)∨(aN∧¬Mb)∨(aP∧bP)；返回 (o_pos∨o_neg, o_neg)")])
_fmla([_t("初始化：h ← IV"), _sub(_t("F"), _t("")), _t("；"), _sub(_t("h"), _t("0")), _t(" ← "), _sub(_t("h"), _t("0")), _t(" ⊕ 0x01010000 ⊕ 32；计数器 64 位（"), _sub(_t("t"), _t("lo")), _t(" / "), _sub(_t("t"), _t("hi")), _t("）")])
_fmla([_t("块压缩（64 字节/块，12 轮）：v ← 双轨载入（轨 A←h、轨 B←rot(h)‖IV；v0⊕=IV4…v6⊕=IV7；v1⊕=t_lo；v9⊕=t_hi；末块 v3⊕=0xFFFFFFFF）")])
_fmla([_t("轮 r=0…11：rA=(3r+skey&7)&31；rB=(7r+(skey>>3)&7)&31；mA=(5r+(skey>>6)&7)&31；mB=(11r+(skey>>9)&7)&31")])
_fmla([_t("轨 A：s0=tadd2(v0,v1,rrp(v2,rA),rrp(v3,rA))；s1=tadd2(v2,v3,rrp(v4,rA+5),rrp(v5,rA+5))")])
_fmla([_t("s2=tadd2(v4,v5,rrp(v6,rA+9),rrp(v7,rA+9))；s3=tadd2(v6,v7,rrp(v0,rA+13),rrp(v1,rA+13))")])
_fmla([_t("p1=tmul2(s0,s2)；p2=tmul2(s1,s3)；majA=quant3(s0,s1,s2)；v2=s1⊕p1；v4=s2⊕p2；v6=s3⊕majA；v7=s3.n⊕rrp(majA,rA+7)")])
_fmla([_t("轨 B 同构（rB/rB+4/rB+8/rB+12、q1/q2/majB，v14=s3b⊕majB；v15=e3.n⊕rrp(majB,rB+6)）")])
_fmla([_t("消息注入：(v0,v1)←tadd2(v0,v1,rrp(M0,mA),rrp(N0,mA))；(v8,v9)←tadd2(v8,v9,rrp(M1,mB),rrp(N1,mB))")])
_fmla([_t("轨间耦合：(v4,v5)←tadd2(v4,v5,v10,v11)；(v12,v13)←tmul2(v12,v13,v0,v1)；(v14,v15)←tadd2(v14,v15,rrp(v6,3r),rrp(v7,3r))")])
_fmla([_t("轮常量错位注入：v0⊕=RCON"), _sub(_t("F"), _t("")), _t("[r&15])；v9⊕=rrp(RCON"), _sub(_t("F"), _t("")), _t("[(r+1)&15], 5r)")])
_fmla([_sub(_t("h"), _t("i")), _t(" ← "), _sub(_t("h"), _t("i")), _t(" ⊕ "), _sub(_t("v"), _t("2i")), _t(" ⊕ rrp("), _sub(_t("v"), _t("2i+1")), _t(", 3i)（i=0…7）；末尾 fin_synth(h)；摘要 ← hex8(h0)‖…‖hex8(h7)，64 hex = 256 位")])

heading("3.4　tsha1b：复杂模型（三进制三轨并行，R=14）", size=12)
para("tsha1b 是复杂档，在 3.2 节双轨骨架基础上**增第三轨海绵轨**，构造三轨并行（R_B=14）："
     "轨 A 与轨 B 复用 tsha1f 的三进制双轨并行扩散（tadd2 平衡加 / tmul2 平衡乘 / quant3 "
     "majority / rrp 旋转混洗，全位运算），另设轨 C 海绵轨 v[16..23]（rate=v[16..19]、"
     "capacity=v[20..23]），每轮对 rate 做消息三进制平面吸收（旋转后平衡加）、8 字环式三进制"
     "置换与独立海绵轮常量注入，随后与轨 A/B 双向耦合；末块把海绵 capacity 折回链值。常量族由"
     "独立种子派生：SEED_B=“TSHA1-2026-b-256-v1”→ IV_B/RCON_B，海绵轨用独立常量 RCON_S "
     "（SEED_S=“TSHA1-2026-b-sponge-v1”）。σ 置换与 S 盒不再需要（该合并显著降低了实现与验证"
     "维护面）。其面向签名、密钥绑定与审计链根等高强度场景提供纵深；输出 8 字大端 hex = 64 hex"
     "（内部摘要 32 字节），整条链压缩后亦经 S=4 最后综合 fin_synth 终筛。")
_fhead("tsha1b 逐步计算式")
_fmla([_t("tadd2 / tmul2 / quant3 定义（见 §A.3.1–A.3.2）；轨 A/轨 B round 结构同 3.3 节，仅轮数取 R_B=14、常量取 IV_B / RCON_B")])
_fmla([_t("块压缩（64 字节/块，14 轮）：工作状态 v[0..23]（双轨 v[0..15] + 海绵轨 v[16..23]）")])
_fmla([_t("海绵轨 v[16..23]：rate=v[16..19]、capacity=v[20..23]；末块 v[19]⊕=0xFFFFFFFF（海绵 padding 边界）")])
_fmla([_t("轮 r=0…13：轨 A/轨 B 三进制扩散与消息注入（同 3.2）+ 海绵吸收/8 字置换/独立常量注入 + 三轨双向耦合 + 轨 A/B 轮常量错位注入")])
_fmla([_t("h_i ← h_i ⊕ v_{2i} ⊕ rrp(v_{2i+1}, 3i) ⊕ rrp(v_{16+2(i&3)}, 7i) ⊕ rrp(v_{17+2(i&3)}, 7i)（三轨折回）；末尾 fin_synth(h)；输出 64 hex = 256 位")])

heading("3.5　tsha1x：加强模型（四轨并行，R=16）", size=12)
para("tsha1x 从 f+b 排列组合范式独立重构为**四轨并行 + 最后综合**（R_X=16，四模型最强）："
     "工作状态 v[0..31]。轨 A（v[0..7]）与轨 B（v[8..15]）复用 trit 位平面双轨扩散（同 f/b "
     "同构，常量用独立 x 常量族），轨 C（v[16..23]）为海绵轨（rate/capacity，复用 §3.4 海绵"
     "语义与独立常量 RCON_S），轨 D（v[24..27]）为 LFSR 反馈移位轨——每轮抽取高位反馈 "
     "fb=(v27>>16)⊕v25，将消息平面旋转后异或吸收进 v24/v25，经轨 A/B/C 字对的 quant3 门控"
     "去线性化后移位（d0←d1、d1←d2、d2←d3、d3←fb⊕qx），再与轨 A/B 双向平衡加混写；四轨各自"
     "注入独立轮常量后进入下一轮。四轨独立、异源的扩散通道（trit 位平面、海绵置换、LFSR 反馈）"
     "在并联中相互交错，显著抬高将单结构分解的代数/差分攻击成本。末段四轨状态投影折叠回链 h 后"
     "执行 fin_synth_x（S=4 一体化收束轮）终筛，输出 8 字大端 hex = 64 hex（内部摘要 32 字节）。"
     "常量族：SEED_X=“TSHA1-2026-x-256-v1”→ IV_X/RCON_X（轨 A/B）；SEED_XD=“TSHA1-2026-x-lfsr-v1”→ "
     "RCON_XD（轨 D）；轨 C 复用 SEED_S 的 RCON_S。")
_fhead("tsha1x 逐步计算式")
_fmla([_t("块压缩（64 字节/块，16 轮）：工作状态 v[0..31]（双轨 v[0..15] + 海绵 v[16..23] + LFSR v[24..27]）")])
_fmla([_t("轨 D 初始化：v24=h2⊕IV0；v25=h3⊕rr32(IV4,11)；v26=h6⊕IV2；v27=h7⊕rr32(IV6,19)")])
_fmla([_t("轮 r=0…15：轨 A/轨 B 三进制扩散与消息注入 → 轨 C 海绵吸收/置换/独立常量注入并与 A/B 双向耦合")])
_fmla([_t("轨 D：sX=(11r+((skey>>18)&7))&31；fb=(v27>>16)⊕v25；v24⊕=rrp(M0,sX)；v25⊕=rrp(N0,sX)；qx=quant3(轨 A/B/C 字对)；d0←d1、d1←d2、d2←d3、d3←fb⊕qx")])
_fmla([_t("轨 D 与轨 A/B 四向平衡加混写；轨 D⊕=RCON_XD[r&15]、v27⊕=rrp(RCON_XD[(r+4)&15],5r)；轨 A/B⊕=RCON_X[r&15]、v9⊕=rrp(RCON_X[(r+1)&15],5r)")])
_fmla([_t("四轨投影折回 h：h_i ← h_i ⊕ v_{2i} ⊕ rrp(v_{2i+1},3i) ⊕ rrp(v_{16+2(i&3)},7i) ⊕ rrp(v_{17+2(i&3)},7i) ⊕ rrp(v_{24+2(i&1)},13i) ⊕ rrp(v_{25+2(i&1)},13i)")])
_fmla([_t("末块海绵 capacity 折回链值后执行 fin_synth_x(h)（S=4 终筛）；输出 64 hex = 256 位")])

heading("3.6　tsha1r：轻量模型（三进制双轨，R=8，128 位）", size=12)
para("tsha1r 面向嵌入式/受限环境，与 f 同为**三进制双轨并行内核**，但轮数最少 R_R=8，"
     "常量族 IV_R/RCON_R 由 SEED_R=“TSHA1-2026-r-128-v1”派生。其整条链压缩后同样复用 "
     "fin_synth 最后综合终筛，但仅取前 4 字输出（32 hex，内部摘要 16 字节）。轮数最少使单次"
     "固定开销为四模型中最低，适合海量小对象的快速指纹（16 字节 → Base48 23 字符）。")
_fhead("tsha1r 逐步计算式")
_fmla([_t("结构完全同 3.3 节（三进制双轨并行：轨 A/轨 B 扩散 + 消息注入 + 轨间耦合 + 轮常量错位注入）")])
_fmla([_t("初始化：h ← IV"), _sub(_t("R"), _t("")), _t("；"), _sub(_t("h"), _t("0")), _t(" ← "), _sub(_t("h"), _t("0")), _t(" ⊕ 0x01010000 ⊕ 16（128 位参数绑定位）")])
_fmla([_t("块压缩（64 字节/块，8 轮）：同 3.3 节但 X while 迭代 R_R=8 轮、常量取 IV_R / RCON_R")])
_fmla([_t("末尾 fin_synth(h)；摘要 ← hex8(h0)‖hex8(h1)‖hex8(h2)‖hex8(h3)，仅前 4 字 = 32 hex = 128 位")])

heading("3.7　常量派生", size=12)
para("为避免任何猜测常量与实现后门争议，全部常量按“标准种子＋PRNG 扩展固化”方式生成："
     "扩展流为 SHA-256(SEED ‖ u64be(k))（k=0,1,2,…）计数器流，依序取字节填充各模型常量。"
     "基础均为 IV（8 个 32 位字）+ 轮常量（16 个 32 位字）：SEED_F=“TSHA1-2026-f-256-v1”、"
     "SEED_R=“TSHA1-2026-r-128-v1”、SEED_B=“TSHA1-2026-b-256-v1”（b 轨 A/B）；海绵轨 C 用独立"
     "常量 RCON_S（SEED_S=“TSHA1-2026-b-sponge-v1”）；x 轨 A/B 用 SEED_X=“TSHA1-2026-x-256-v1”、"
     "轨 D 用 SEED_XD=“TSHA1-2026-x-lfsr-v1”。σ 置换与 S 盒不再需要。"
     "该过程可复现（tests/tsha_probe/gen_tsha1fr.py、gen_tsha1bx.py 与 gen_tsha1x.py），"
     "且 IV 与 BLAKE2 的 6a09e667… 有意区分。")

heading("3.8　输出编码与位长机制", size=12)
para("默认输出使用自行设计的 Base48 编码（std/base48.tie，命名空间 b48）：字符台为 48 符号"
     "连续集合“0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKL”（10 数字＋26 小写＋12 "
     "大写，顺序即数值 0..47）。n 字节（8n 位）到 m 个 Base48 字符的换算为 m=ceil(8n/"
     "log2(48))，取整采用精确有理近似 m=(n×1432407)/1000000+1；编码实现为 base-256 字节数组"
     "的反复长除 48（进位中间量 <2^13，无溢出），解码为反复“×48+digit”（同样无溢出），并采用"
     "长度保留设计使 encode/decode 互为精确逆（前导零字节完整往返）。")
para("位长机制（输出层）：内部压缩固定 32/16 字节摘要，位长在输出层调整、以 n 个 48 进制"
     "符号为准。基础块 = 内部摘要的 Base48 编码（32 字节 → 46 字符、16 字节 → 23 字符）；若"
     "基础块符号数 ≥ n 直接取前 n 个符号；不足（n=64/88/96 超过 46/23）则按 XOF 流追加扩展块"
     "——第 k 块 = digest(摘要 ‖ u64be64(k))（32 字节，k=0,1,2,…）编码 Base48 拼接，直到符号数"
     "≥ n 再取前 n 个符号，块级拼接保证同模型 n 较大的输出以较小位长输出为前缀。对外仅允许 "
     "{2,3,4,6,8,12,16,24,32,48,64,88,96}；其它进制由 n 先换算内嵌字节数 b0=⌊n·log2(48)/8⌋，"
     "池取前 b0 字节后按 base 编码。同时保留 _hex 变体（tsha1f_hex/b_hex/x_hex 为 64 hex、"
     "tsha1r_hex 为 32 hex），满足十六进制生态需求。字符台用户选定，集合内含 0↔o、1↔l/I "
     "视觉相近对，属换取字符连续性的取舍，在实现中明确注释说明。")

# ================= 4 安全性质 =================
heading("4　安全性质与论证")
para("TSHA1“更安全”的主张建立在四类可论证机制上，而非经验式声明：(a) 结构层面——"
     "抗长度扩展（计数器式纯零填充＋64 位长度绑定，末块全 1 标记，§3.2）与抗冲击面收窄"
     "（状态与输出解耦）；(b) 构造层面——三进制平衡加/乘与 majority 三类非线性经 trit 位平面"
     "化后以错位移位在双轨内并行扩散、并以轨间耦合与轮常量错位注入混合，多个异源非线性通道"
     "交织，显著抬高差分/代数类结构化攻击所需的多轮追踪复杂度；加强模型 x 的异源四轨并行"
     "（§3.5）进一步将攻击者对单结构的利用性切割；(c) 信任层面——自创层严格限于参数/"
     "常量/轮数与并行结构，f/b/r 共享统一压缩原语使实现可验证面最小，未引入任何未经审计的全新"
     "结构断言，安全性可随组合逐一论证；(d) 工程层面——常量种子化可复现（§3.7）、实现纯 "
     "TIE 可自举交付（§5.1）、KAT 逐字节核对（§5.2），从源头排除后门注入与实现漂移。")
para("以下逐点说明安全边界与已知限制：(1) 结构信任边界——自创层未宣称新结构贡献，"
     "属确定性组合，家族在说明文档中以 security-notes 声明**未经独立审计**，正式部署前"
     "应由第三方密码学审计复核；(2) 雪崩与差分——tadd2/tmul2/quant3 提供逐 trit 平衡三进制"
     "扩散，向量探针判定实现与生成器一致，但一致性验证不等同于统计雪崩检验，后续将补做"
     "严格的差分/雪崩统计测试；(3) 长度扩展免疫——计数器式纯零填充与末块标记在结构上杜绝 "
     "Merkle–Damgård 型长度扩展（四模型统一）；(4) 常数时间——TIE 数值执行语义不提供硬件级"
     "常数时间保证（无 volatile、无恒定时间指令特判），位平面打包与 24 基调度引入的分支循环"
     "在原理上允许时间侧信道差异；TSHA1 当前用于完整性指纹与签名对象哈希，不直接处理机密"
     "数据；若后续用于常时间敏感场景，须在平台层落实常数时间保证并复核；(5) 输出编码——"
     "Base48 默认输出（§3.8）不改变内部 256/128 位摘要字节，decode(b48) 与 hex 逐字节一致。")

# ================= 5 实现与评估 =================
heading("5　实现与评估")

heading("5.1　纯 TIE 实现", size=12)
para("全部算法以 TIE 编写（std/tsha1.tie、std/base48.tie，命名空间 tsha/b48），仅依赖"
     "语言底座原语（str_len/str_char/table 动态表/i64 位运算），零外部运行时依赖；字符串"
     "按 {ptr,len} 二进制安全模型逐字节读写，支持含 NUL 在内的任意字节输入。f/r 复用同一份"
     "三进制双轨压缩内核（compress_f/compress_r 同构）+ 共享 fin_synth，compress_b 在双轨上"
     "增海绵第三轨，x 用 compress_x 四轨并行与 fin_synth_x；仅轮数与常量区分，代码复用度高。"
     "验证探针（tests/tsha_probe/，含 gen_tsha1fr.py、gen_tsha1bx.py、gen_tsha1x.py）以编译-运行"
     "方式断言四模型全部 KAT 与 hex↔Base48 往返一致，编译零错误；探针独立于常规回归（regress-s21）运行。")

heading("5.2　已知答案向量（KAT）", size=12)
add_table(
    ["输入", "tsha1f（64 hex）", "tsha1r（32 hex）", "tsha1b（64 hex）", "tsha1x（64 hex）"],
    [
        ["空串 \"\"", "45f2029d 37ad41cf 835735aa 3f0c1b64 a2980fe2 64f83d35 08d10119 64713550",
         "35a4dc97 79e6e2e6 5b05790f 491a533f",
         "afd3b3fc bb903a5a 55f7ce0e 79b83bf1 5edfefbc 495eaec0 82ce9dce 71afe1f8",
         "157daf00 76845caa 9b1bae59 c40765dd 9dc8bab0 49d59cec 7a1342c6 39fc5d96"],
        ["\"abc\"", "f38e5ae5 a532e478 3578d24c 7ade7a9c fcb1f384 9a295c99 92043c17 9d8d26c0",
         "314db448 edf50a42 64d29b00 bdedca6f",
         "4bbaea17 4cd4475b 2b17e052 fa3f9f52 3c2cd7c8 10025837 2311bd72 88aa770a",
         "8baa7531 12e9a18f 8393f912 745d2d40 a17a750c 34d5e970 f6b28484 b53a543b"],
        ["1000×'a'", "27bce163 51ee687d 56541ea8 b0339f01 b9c76931 c5765098 2488e331 ceb54ebd",
         "19f291e4 1241e4bb 89d0c95d 06e2b810",
         "1e9e0841 7fce2c02 189d439d 7ef25aaf d81465d8 7c18ce44 411840ca ee8d45d5",
         "be64d8ac b64239cf 11b5a13e a3bbf1df 16baa4e1 ee3f5690 ba124cd3 569c2bfa"],
    ],
    "表 2　TSHA1 已知答案向量（KAT，与平台生成器交叉核对，内部摘要 hex）")

heading("5.3　性能基准", size=12)
para("基准方法：计时原语为秒级粒度，故采用“目标时长循环”——对每个（模型，消息长度）"
     "组合重复哈希至累计≥2 s，以总数据量/耗时折算吞吐；消息构造使用 StringBuilder 原地"
     "追加（规避字符串 + 的 O(n²) 拼接陷阱——该陷阱与治理详见语言文档）。环境为 TIE 自写"
     "编译器前端驱动、LLVM 后端生成 x86_64 可执行文件的执行层。实测大消息稳态吞吐与短消息"
     "哈希率见表 3。")
add_table(
    ["模型", "2 KB（hash/s）", "2 KB（MB/s）", "32 KB（hash/s）", "512 KB（hash/s）", "512 KB（MB/s）"],
    [
        ["tsha1f", "3804", "7", "909", "66", "33"],
        ["tsha1b", "3646", "7", "498", "33", "16"],
        ["tsha1x", "2992", "5", "332", "23", "11"],
        ["tsha1r", "6732", "13", "1200", "84", "42"],
    ],
    "表 3　TSHA1 四模型基准实测（目标时长 2s 采样，短消息模型含秒级量化误差）")
para("观察：(1) 大消息稳态吞吐排序 r≈42 > f≈33 > b≈16 > x≈11 MB/s，与设计模型一致——"
     "轻量档轮数最少（R_R=8）最快，快速档次之，复杂档因三轨、加强档因四轨（R_X=16）而更慢，"
     "具更高的抗破解冗余；(2) 短消息（2 KB）单次固定开销排序为 r（6732 hash/s）> f（3804）"
     "> b（3646）> x（2992），轻量档仍最优——因此海量小文件指纹场景宜选 r，需要更高强度时可换 b；"
     "(3) 绝对吞吐受 TIE 自写编译器/IR 执行层的软件实现约束（无 SIMD），较原生 C（BLAKE2s "
     "可达 GB/s 级）低 2~3 个数量级，但对审计链 per-file 指纹与低频签名对象场景完全充足。")

# ================= 6 应用 =================
heading("6　应用：插件化内核审计链")
para("TIE 全平台插件化内核以“核心微内核（机制）+注册表+审计器+加载器”为骨架，所加载"
     "插件（包）须通过八级审计链：①指纹树重算→②包内验签→③公钥指纹锚定→④IR 版本→"
     "⑤id/version→⑥字段白名单→⑦依赖解析→⑧注册仲裁。其中指纹模型采用 TSHA1："
     "单文件级指纹用快速模型 tsha1f，整包树根（即签名对象）用加强模型 tsha1x；凭证区的"
     " Ed25519 纯 TIE 实现（RFC 8032 向量绿）对 tsha1x 树根签名。该分层与基准实测吻合："
     "文件级批量校验取最实惠的快速/轻量模型，跨包完整性/冒名抵御取最高强度模型，二者速度差约 "
     "3× 构成清晰的成本梯度。")

# ================= 7 结论 =================
heading("7　结论与展望")
para("本文提出了 TSHA1——一个完全以自举语言 TIE 实现的**新型安全杂凑函数家族**。其"
     "新颖性在于统一的平衡三进制位平面压缩原语与双轨骨架：f/r 复用同一套 trit 位平面双轨并行"
     "压缩器（平衡加/平衡乘/majority，全位运算）与最后综合，b 增第三轨海绵轨，x 独立重构为"
     "四轨并行（双轨 + 海绵 + LFSR，R_X=16），仅轮数与种子派生常量/并行结构不同，从而以最小"
     "实现面构建四档速度/强度梯度。更强的安全性质"
     "来自四类可论证机制——结构层面抗长度扩展、构造层面多非线性构件交叠、信任层面自创层"
     "限于参数/常量/轮数与并行结构、工程层面常量可复现与纯 TIE 自举交付。实现为纯 TIE、零外部"
     "依赖，探针逐字节核对 KAT、字节边界全覆盖；基准结果与模型设计一致，并已投入插件化内核"
     "的指纹与凭证签名底座。后续工作包括：严格的差分/雪崩统计检验与独立第三方安全审计；将"
     "位平面扩散与既有 Keccak/Ascon 置换统一为可配置扩散内核以降低维护面；在编译器执行层"
     "实现字符串就地追加等语言级性能优化后重新量化吞吐；以及将 TSHA1 与后量子签名（SLH-DSA "
     "纯 TIE 前沿）组合为审计链的长期凭证方案。")

# ================= 参考文献 =================
heading("参考文献")
refs = [
    "NIST. FIPS 180-4: Secure Hash Standard (SHA-2)[S]. Gaithersburg: NIST, 2015.",
    "NIST. FIPS 202: SHA-3 Standard: Permutation-Based Hash and Extendable-Output Functions[S]. Gaithersburg: NIST, 2015.",
    "Saarinen M J O, Aumasson J P. RFC 7693: The BLAKE2 Cryptographic Hash and Message Authentication Code (MAC)[S]. IETF, 2015.",
    "Aumasson J P, Neves S, Wilcox-O'Hearn Z, Winnerlein C. BLAKE2: Simpler, Smaller, Fast as MD5[C]// ACNS 2013, LNCS 7954. Springer, 2013: 119-135.",
    "O'Connor J P, Aumasson J P, Neves S, Wilcox-O'Hearn Z. BLAKE3: One Function, Fast Everywhere[R/OL]. 2020. https://github.com/BLAKE3-team/BLAKE3.",
    "Dobraunig C, Eichlseder M, Mendel F, Schläffer M. Ascon v1.2: Lightweight Authenticated Encryption and Hashing[R]. NIST Lightweight Cryptography Finalist, 2021.",
    "NIST. FIPS 203: Module-Lattice-Based Key-Encapsulation Mechanism Standard (ML-KEM)[S]. Gaithersburg: NIST, 2024.",
    "NIST. FIPS 205: Stateless Hash-Based Digital Signature Standard (SLH-DSA)[S]. Gaithersburg: NIST, 2024.",
    "Bernstein D J. Curve25519: New Diffie-Hellman Speed Records[C]// PKC 2006, LNCS 3958. Springer, 2006: 207-228.",
    "Josefsson S, Liusvaara I. RFC 8032: Edwards-Curve Digital Signature Algorithm (EdDSA)[S]. IETF, 2017.",
    "Bertoni G, Daemen J, Peeters M, Van Assche G. The Keccak Reference[R]. Keccak Team, 2011.",
    "NIST. FIPS 204: Module-Lattice-Based Digital Signature Standard (ML-DSA)[S]. Gaithersburg: NIST, 2024.",
    "Krawczyk H, Bellare M, Canetti R. RFC 2104: HMAC: Keyed-Hashing for Message Authentication[S]. IETF, 1997.",
    "TIE Language Project. TIE: A Self-Hosting Programming Language — AGENTS.md, Architecture Notes and Design Specs[R/OL]. 2026. tie-lang.",
]
for i, r in enumerate(refs):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.9)
    pf.first_line_indent = Cm(-0.9)   # 悬挂缩进
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    run = p.add_run("[%d] %s" % (i + 1, r))
    _set_font(run, "宋体", "Times New Roman", 10.5)

# ================= 附录：统一内核逐步计算公式 =================
heading("附录　统一内核逐步计算公式（与 std/tsha1.tie 逐式对应）", size=14)
para("本附录按步给出统一内核（trit 位平面原语与双轨骨架）的逐步计算公式，与标准库实现"
     "（std/tsha1.tie）逐式对应；b 在该双轨基础上增海绵第三轨、x 为四轨并行，其增补公式见 "
     "A.3/A.4。记号约定：⊕ 表示按位异或，∧、∨、¬ 分别表示按位与、或、非；"
     "ROR(n,x) 与 ROL(n,x) 表示对 32 位字 x 循环右移 / 左移 n 位；rrp(x,r)=ROR(r&31,x)；"
     "hex8(w) 把字 w 按大端字节序写成 8 个 hex 字符；A‖B 表示串接。字节 → 调度基础量：字节 b "
     "高两位 tf=(b>>6)&3 经饱和映射 tv(b)∈{−1,0,+1}（0→−1、1→0、2/3→+1），低三位 "
     "low3(b)=b&7，合成 24 基数字 d(b)=3·low3(b)+(tv(b)+1)∈[0,23]。trit 位平面：每 32 个 trit "
     "打包为幅值平面 M（第 i 位=1 ⟺ 第 i 个 trit 非零）与符号平面 N（第 i 位=1 ⟺ 第 i 个 trit "
     "为负），记 (M,N)=Pack(B)。消息 64 字节块派生两组位平面 (M0,N0),(M1,N1) 与 24 基调度键 "
     "skey ← mod(∑_j d(B[j])·24^(63−j), 2^32)。")
para("**全部常量（每模型 IV 8 字 + 轮常量 16 字作基础）由标准种子 + PRNG 扩展（SHA-256(SEED ‖ "
     "u64be(k)) 计数器流）确定性派生（§3.7）：IV_F/RCON_F、IV_R/RCON_R、IV_B/RCON_B，海绵轨 C 的 "
     "RCON_S，x 轨 A/B 的 IV_X/RCON_X 与轨 D 的 RCON_XD；具体十六进制值见常量生成器（tests/tsha_probe/"
     "gen_tsha1fr.py、gen_tsha1bx.py 与 gen_tsha1x.py）。**")

heading("A.1　统一双轨骨架 compress(h, B, t, last)（f/r 共用，仅 R 与常量不同）", size=12)
_fhead("A.1.1　双轨载入与块初始化")
_fmla([_t("1. 消息平面与 24 基调度键：(M0,N0),(M1,N1) ← Pack(B)（64 字节 → 两组 32-trit 平面）；skey ← mod(∑_j d(B[j])·24^(63−j), 2^32)")])
_fmla([_t("2. 工作状态 v[0..15]（轨 A=v[0..7]、轨 B=v[8..15]），从链值 h 与 IV 载入：")])
_fmla([_t("　 v0=h0, v1=ROR(7,h4), v2=h1, v3=ROR(7,h5), v4=h2, v5=ROR(7,h6), v6=h3, v7=ROR(7,h7)")])
_fmla([_t("　 v8=ROR(13,h0), v9=IV0, v10=ROR(13,h1), v11=IV1, v12=ROR(13,h2), v13=IV2, v14=ROR(13,h3), v15=IV3")])
_fmla([_t("3. 分块注入：v0⊕=IV4；v2⊕=IV5；v4⊕=IV6；v6⊕=IV7；v1⊕=t_lo；v9⊕=t_hi；若 last 则 v3⊕=0xFFFFFFFF")])
_fhead("A.1.2　逐轮三进制双轨并行（R 轮，f=12、r=8）")
_fmla([_t("对 r = 0,1,…,R−1（R 见模型）")])
_fmla([_t("　 a. 移位量：rA=(3r+(skey&7))&31；rB=(7r+((skey>>3)&7))&31；mA=(5r+((skey>>6)&7))&31；mB=(11r+((skey>>9)&7))&31")])
_fmla([_t("　 b. 轨 A 三进制扩散（balanced add + balanced mul + majority）：")])
_fmla([_t("　　　s0=tadd2(v0,v1,rrp(v2,rA),rrp(v3,rA))；s1=tadd2(v2,v3,rrp(v4,rA+5),rrp(v5,rA+5))")])
_fmla([_t("　　　s2=tadd2(v4,v5,rrp(v6,rA+9),rrp(v7,rA+9))；s3=tadd2(v6,v7,rrp(v0,rA+13),rrp(v1,rA+13))")])
_fmla([_t("　　　p1=tmul2(s0,s2)；p2=tmul2(s1,s3)；majA=quant3(s0,s1,s2)")])
_fmla([_t("　　　v0=s0.M, v1=s0.N；v2=s1.M⊕p1.M, v3=s1.N⊕p1.N；v4=s2.M⊕p2.M, v5=s2.N⊕p2.N；v6=s3.M⊕majA, v7=s3.N⊕rrp(majA,rA+7)")])
_fmla([_t("　 c. 轨 B 同构（移位量 rB、rB+4、rB+8、rB+12；q1=tmul2(u0,u2)、q2=tmul2(u1,u3)、majB=quant3(u0,u1,u2)）：")])
_fmla([_t("　　　v8=u0.M, v9=u0.N；v10=u1.M⊕q1.M, v11=u1.N⊕q1.N；v12=u2.M⊕q2.M, v13=u2.N⊕q2.N；v14=u3.M⊕majB, v15=u3.N⊕rrp(majB,rB+6)")])
_fmla([_t("　 d. 消息平面注入（旋转后平衡加）：(v0,v1)←tadd2(v0,v1,rrp(M0,mA),rrp(N0,mA))；(v8,v9)←tadd2(v8,v9,rrp(M1,mB),rrp(N1,mB))")])
_fmla([_t("　 e. 轨间耦合：(v4,v5)←tadd2(v4,v5,v10,v11)；(v12,v13)←tmul2(v12,v13,v0,v1)；(v14,v15)←tadd2(v14,v15,rrp(v6,3r),rrp(v7,3r))")])
_fmla([_t("　 f. 轮常量双轨错位注入：v0⊕=RCON[r&15]；v9⊕=rrp(RCON[(r+1)&15], 5r)")])
_fhead("A.1.3　双轨折回与最后综合")
_fmla([_t("4. 双轨折回链值：h_i ← h_i ⊕ v_{2i} ⊕ rrp(v_{2i+1}, 3i)（i=0,…,7）")])
_fmla([_t("5. 最后综合 fin_synth(h)：把 h 重新载入双轨（v0..v7=h0..h7；v8..v15=ROR(13,h0..h3)、ROR(13,h4..h7)）")])
_fmla([_t("6. 对 sf=0..3（S=4 轮），对 t=0..7：t1=(t+1)&7、t2=(t+2)&7、t3=(t+3)&7、t4=(t+4)&7；rot=(3sf+5t)&31")])
_fmla([_t("　 am=tadd2(v_{2t},v_{2t+1},rrp(v_{2t1},rot),rrp(v_{2t1+1},rot))")])
_fmla([_t("　 pm=tmul2(v_{2t3},v_{2t3+1},rrp(v_{2t4},rot+1),rrp(v_{2t4+1},rot+1))；mj=quant3(v_{2t},v_{2t+1},v_{2t2},v_{2t2+1},am.M,am.N)")])
_fmla([_t("　 tmp_{2t}=am.M⊕pm.M⊕mj；tmp_{2t+1}=am.N⊕pm.N⊕rrp(mj,rot+3)，整轮赋值回 v")])
_fmla([_t("7. 投影折叠：h_i ← h_i ⊕ v_{2i} ⊕ ROL((7i)&31, v_{2i+1})（i=0,…,7）")])

heading("A.2　平衡三进制原语 tadd2 / tmul2 / quant3", size=12)
_fhead("A.2.1　tadd2((Ma,Na),(Mb,Nb)) → (Mo,No)（平衡三进制两位 trit 相加，位切片）")
_fmla([_t("aP=Ma∧¬Na；aN=Ma∧Na；bP=Mb∧¬Nb；bN=Mb∧Nb")])
_fmla([_t("o_pos=(¬Ma∧bP)∨(aP∧¬Mb)∨(aN∧bN)")])
_fmla([_t("o_neg=(¬Ma∧bN)∨(aN∧¬Mb)∨(aP∧bP)")])
_fmla([_t("Mo = o_pos ∨ o_neg；No = o_neg")])
_fhead("A.2.2　tmul2((Ma,Na),(Mb,Nb)) → (Mo,No)（±1·±1=±1，0 吸收）")
_fmla([_t("Mo = Ma ∧ Mb（两 trit 均非零才非零）")])
_fmla([_t("No = (Na ⊕ Nb) ∧ Mo（± 相异才为负）")])
_fhead("A.2.3　quant3((M0,N0),(M1,N1),(M2,N2))（majority 量化）")
_fmla([_t("A = M0 ∧ ¬N0；B = M1 ∧ ¬N1；C = M2 ∧ ¬N2")])
_fmla([_t("输出 = (A ∧ B) ∨ (B ∧ C) ∨ (C ∧ A)（三路 majority，+1 计数 ≥ 2 才为 1，含平局→0）")])

heading("A.3　digest_f / digest_b / digest_r 全流程（同为计数器式纯零填充）", size=12)
_fhead("A.3.1　digest_f(msg)（R=12，常量 IV_F/RCON_F）")
_fmla([_t("1. h ← IVf；h0 ← h0 ⊕ 0x01010000 ⊕ 32（256 位参数绑定位）；t ← 0；")])
_fmla([_t("2. 按 64 字节切块；每个完整块：t 累加（t_hi 随进位），compress(h, B, t, false)；")])
_fmla([_t("3. 末块（不足 64 字节，以 0 填充）：t ← t + rem，compress(h, B, t, true)；")])
_fmla([_t("4. fin_synth(h)；摘要 ← hex8(h0)‖…‖hex8(h7)（64 hex 字符 = 256 位）")])
_fhead("A.3.2　digest_b(msg)（R=14，常量 IV_B/RCON_B + 海绵轨 RCON_S）")
_fmla([_t("1. h ← IVb；h0 ← h0 ⊕ 0x01010000 ⊕ 32（256 位参数绑定位）；sp[0..3]←0；t ← 0；")])
_fmla([_t("2. 64 字节块迭代：t 累加，compress_b(h, B, t, false, sp)；末块 compress_b(h, B, t, true, sp)；")])
_fmla([_t("3. compress_b 内：双轨 v[0..15] + 海绵轨 v[16..23]（rate=v[16..19]、capacity=v[20..23]）；14 轮三轨并行；末块 capacity 填回 sp")])
_fmla([_t("4. 末块海绵 capacity 折回链值：h0⊕=sp0；h1⊕=rr32(sp1,7)；h2⊕=sp2；h3⊕=rr32(sp3,7)")])
_fmla([_t("5. fin_synth(h)；摘要 ← hex8(h0)‖…‖hex8(h7)（64 hex 字符 = 256 位）")])
_fhead("A.3.3　digest_r(msg)（R=8，常量 IV_R/RCON_R，128 位）")
_fmla([_t("1. h ← IVr；h0 ← h0 ⊕ 0x01010000 ⊕ 16（128 位参数绑定位）；t ← 0；")])
_fmla([_t("2–3. 同 A.3.1 的 64 字节块迭代与末块（纯零填充 + 计数器 + 末块全 1），但 while 迭代 8 轮、常量取 IV_R/RCON_R")])
_fmla([_t("4. fin_synth(h)；摘要 ← hex8(h0)‖hex8(h1)‖hex8(h2)‖hex8(h3)（仅前 4 字 = 32 hex = 128 位）")])

heading("A.4　tsha1x（加强模型，256 位，四轨并行含 LFSR，R_X=16）", size=12)
_fhead("A.4.1　digest_x(msg) 全流程（R_X = 16 轮）")
_fmla([_t("1. h ← IVx；h0 ← h0 ⊕ 0x01010000 ⊕ 32；sp[0..3]←0；t ← 0；")])
_fmla([_t("2. 64 字节块迭代（t 累加，compress_x(h, B, t, false, sp)）；末块 compress_x(h, B, t, true, sp)；")])
_fmla([_t("3. compress_x 内工作状态 v[0..31]：填轨 A（v[0..7]）、轨 B（v[8..15]）、轨 C 海绵（v[16..23]，rate=v[16..19]、capacity=v[20..23]）、轨 D LFSR（v[24..27]）")])
_fmla([_t("4. 轨 D 初始化 v24=h2⊕IV0；v25=h3⊕rr32(IV4,11)；v26=h6⊕IV2；v27=h7⊕rr32(IV6,19)")])
_fmla([_t("5. 16 轮四轨并行：轨 A/B 三进制扩散与消息注入 → 轨 C 海绵（吸收/8 字置换/RCON_S/与 A/B 耦合）→ 轨 D（fb 反馈、吸收、quant3 门控移位、与 A/B 四向混写、RCON_XD）→ 轨 A/B⊕RCON_X")])
_fmla([_t("6. 四轨状态投影折回 h：h_i ← h_i ⊕ v_{2i} ⊕ rrp(v_{2i+1},3i) ⊕ rrp(v_{16+2(i&3)},7i) ⊕ rrp(v_{17+2(i&3)},7i) ⊕ rrp(v_{24+2(i&1)},13i) ⊕ rrp(v_{25+2(i&1)},13i)")])
_fmla([_t("7. 末块海绵 capacity 折回链值（同 A.3.2 第 4 步）后 fin_synth_x(h)；摘要 ← hex8(h0)‖…‖hex8(h7)（64 hex = 256 位）")])

heading("A.5　输出编码与位长（全族共享）", size=12)
_fmla([_t("任选一模型：内部十六进制摘要 H（32/16 字节）→ Base48 编码（b48.encode，见 std/base48.tie）：")])
_fmla([_t("m = ⌈8L / log2 48⌉ = (L·1432407)/10^6 + 1（L=len(H)/2 字节 → m 个 48 进制字符）")])
_fmla([_t("位长 n（48 进制符号个数，仅允许 {2,3,4,6,8,12,16,24,32,48,64,88,96}）：若基础块符号数 ≥ n 取前 n 个；")])
_fmla([_t("不足则按 XOF 流扩展：第 k 块 = digest(摘要 ‖ u64be64(k)) 编码 Base48 拼接，直到符号数 ≥ n；")])
_fmla([_t("其它进制 {2,3,8,16}：由 n 换算内嵌字节数 b0=⌊n·log2(48)/8⌋，池取前 b0 字节后按 base 重编码为 m′ 个符号（base16 → 恒等）")])

OUT2 = OUT[:-5] if OUT.lower().endswith(".docx") else OUT
try:
    doc.save(OUT)
    final = OUT
except PermissionError:
    alt = OUT2 + "-v3.docx"
    doc.save(alt)
    final = alt
print("saved:", final)